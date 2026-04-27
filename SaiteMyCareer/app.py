from flask import Flask, render_template, redirect, url_for, request, jsonify, flash, session, send_file
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime, date
import json
import os
from database import db, User, Event, EventRegistration, VolunteerHours, Notification
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here-change-this-in-production'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///youth_platform.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['ALLOWED_EXTENSIONS'] = {'png', 'jpg', 'jpeg', 'gif', 'pdf', 'doc', 'docx'}

db.init_app(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Пожалуйста, войдите для доступа к этой странице'
login_manager.login_message_category = 'info'

# Создание папок для загрузки
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs('images', exist_ok=True)


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def format_phone(phone):
    """Форматирует номер телефона"""
    if not phone:
        return phone
    digits = re.sub(r'\D', '', phone)
    if len(digits) == 11 and digits.startswith('8'):
        digits = '7' + digits[1:]
    if len(digits) == 11:
        return f"+7({digits[1:4]}){digits[4:7]}-{digits[7:9]}-{digits[9:11]}"
    return phone


def normalize_phone_for_db(phone):
    """Нормализует номер телефона для поиска в БД"""
    if not phone:
        return phone
    digits = re.sub(r'\D', '', phone)
    if len(digits) == 11 and digits.startswith('8'):
        digits = '7' + digits[1:]
    return digits


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


@app.context_processor
def utility_processor():
    return {
        'now': datetime.now(),
        'today': date.today()
    }


@app.template_filter('from_json')
def from_json(value):
    if value:
        try:
            return json.loads(value)
        except:
            return []
    return []


@app.route('/')
def index():
    upcoming_events = Event.query.filter(
        Event.event_end_date >= date.today()
    ).order_by(Event.event_start_date).limit(6).all()
    return render_template('index.html', upcoming_events=upcoming_events)


@app.route('/events')
def events():
    events = Event.query.filter(
        Event.event_end_date >= date.today()
    ).order_by(Event.event_start_date).all()
    return render_template('events.html', events=events)


@app.route('/event/<int:event_id>')
def event_detail(event_id):
    event = Event.query.get_or_404(event_id)
    return render_template('event_detail.html', event=event)


@app.route('/about')
def about():
    return render_template('about.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        remember = True if request.form.get('remember') else False

        search_username = normalize_phone_for_db(username)
        users = User.query.all()
        user = None
        for u in users:
            if normalize_phone_for_db(u.username) == search_username:
                user = u
                break

        if user and check_password_hash(user.password, password):
            login_user(user, remember=remember)
            next_page = request.args.get('next')
            flash(f'Добро пожаловать, {user.get_full_name()}!', 'success')
            return redirect(next_page) if next_page else redirect(url_for('index'))
        else:
            flash('Неверный логин или пароль', 'error')

    return render_template('login.html')


@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        username = request.form.get('username')
        secret_word = request.form.get('secret_word')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')

        search_username = normalize_phone_for_db(username)
        users = User.query.all()
        user = None
        for u in users:
            if normalize_phone_for_db(u.username) == search_username:
                user = u
                break

        if not user:
            flash('Пользователь с таким номером не найден', 'error')
            return redirect(url_for('forgot_password'))

        if user.secret_word != secret_word:
            flash('Неверное кодовое слово', 'error')
            return redirect(url_for('forgot_password'))

        if new_password != confirm_password:
            flash('Пароли не совпадают', 'error')
            return redirect(url_for('forgot_password'))

        user.password = generate_password_hash(new_password)
        db.session.commit()

        flash('Пароль успешно изменен! Теперь вы можете войти', 'success')
        return redirect(url_for('login'))

    return render_template('forgot_password.html')


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        password_confirm = request.form.get('password_confirm')
        secret_word = request.form.get('secret_word')
        last_name = request.form.get('last_name')
        first_name = request.form.get('first_name')
        middle_name = request.form.get('middle_name')
        birth_date = request.form.get('birth_date')
        email = request.form.get('email')
        experience = request.form.get('experience')

        formatted_phone = format_phone(username)
        normalized_phone = normalize_phone_for_db(username)

        if not username or not password or not password_confirm or not secret_word:
            flash('Все обязательные поля должны быть заполнены', 'error')
            return redirect(url_for('register'))

        if password != password_confirm:
            flash('Пароли не совпадают', 'error')
            return redirect(url_for('register'))

        if len(password) < 4:
            flash('Пароль должен содержать минимум 4 символа', 'error')
            return redirect(url_for('register'))

        users = User.query.all()
        for u in users:
            if normalize_phone_for_db(u.username) == normalized_phone:
                flash('Пользователь с таким номером уже существует', 'error')
                return redirect(url_for('register'))

        if email and User.query.filter_by(email=email).first():
            flash('Пользователь с таким email уже существует', 'error')
            return redirect(url_for('register'))

        user = User(
            username=formatted_phone,
            password=generate_password_hash(password),
            secret_word=secret_word,
            role='volunteer',
            phone=formatted_phone,
            last_name=last_name,
            first_name=first_name,
            middle_name=middle_name,
            birth_date=datetime.strptime(birth_date, '%Y-%m-%d') if birth_date else None,
            email=email,
            experience=experience
        )

        db.session.add(user)
        db.session.commit()

        flash('Регистрация успешна! Теперь вы можете войти', 'success')
        return redirect(url_for('login'))

    return render_template('register.html')


@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Вы вышли из системы', 'info')
    return redirect(url_for('index'))


@app.route('/profile')
@login_required
def profile():
    registrations = EventRegistration.query.filter_by(user_id=current_user.id).order_by(
        EventRegistration.registered_at.desc()).all()
    hours = VolunteerHours.query.filter_by(user_id=current_user.id).order_by(VolunteerHours.created_at.desc()).all()

    return render_template('profile.html',
                           user=current_user,
                           registrations=registrations,
                           hours=hours)


@app.route('/profile/update', methods=['POST'])
@login_required
def update_profile():
    user = current_user

    last_name = request.form.get('last_name')
    first_name = request.form.get('first_name')
    middle_name = request.form.get('middle_name')
    phone = request.form.get('phone')
    email = request.form.get('email')
    social_links = request.form.get('social_links')
    about = request.form.get('about')
    experience = request.form.get('experience')
    birth_date = request.form.get('birth_date')

    if last_name is not None:
        user.last_name = last_name
    if first_name is not None:
        user.first_name = first_name
    if middle_name is not None:
        user.middle_name = middle_name
    if phone is not None:
        user.phone = phone
    if email is not None:
        user.email = email
    if social_links is not None:
        user.social_links = social_links
    if about is not None:
        user.about = about
    if experience is not None:
        user.experience = experience

    if birth_date:
        try:
            user.birth_date = datetime.strptime(birth_date, '%Y-%m-%d')
        except:
            pass

    avatar_updated = False
    if 'avatar' in request.files:
        file = request.files['avatar']
        if file and file.filename:
            if allowed_file(file.filename):
                file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
                filename = secure_filename(f"user_{user.id}_{datetime.now().timestamp()}.{file_ext}")
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)

                if user.avatar and user.avatar != 'default-avatar.jpg':
                    old_avatar_path = os.path.join(app.config['UPLOAD_FOLDER'], user.avatar)
                    if os.path.exists(old_avatar_path):
                        try:
                            os.remove(old_avatar_path)
                        except:
                            pass

                user.avatar = filename
                avatar_updated = True
                flash('Аватарка успешно обновлена', 'success')
            else:
                flash('Недопустимый формат файла. Разрешены: png, jpg, jpeg, gif', 'error')

    try:
        db.session.commit()
        if not avatar_updated and any(
                [last_name, first_name, middle_name, phone, email, social_links, about, experience, birth_date]):
            flash('Профиль обновлен', 'success')
    except Exception as e:
        db.session.rollback()
        flash('Ошибка при сохранении профиля', 'error')
        print(f"Error updating profile: {e}")

    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'success': True, 'avatar': user.avatar})

    return redirect(url_for('profile'))


@app.route('/event/<int:event_id>/register/<string:reg_type>', methods=['POST'])
def register_for_event(event_id, reg_type):
    try:
        event = Event.query.get_or_404(event_id)

        if not event.is_registration_open():
            return jsonify({'error': 'Регистрация на это мероприятие закрыта'}), 400

        if reg_type not in ['participant', 'volunteer']:
            return jsonify({'error': 'Неверный тип регистрации'}), 400

        if reg_type == 'volunteer' and not event.use_volunteer_template:
            return jsonify({'error': 'Регистрация волонтеров на это мероприятие недоступна'}), 400

        if reg_type == 'participant' and not event.use_participant_template:
            return jsonify({'error': 'Регистрация участников на это мероприятие недоступна'}), 400

        data = request.get_json()
        print(f"Received registration data: {data}")

        # Проверка согласия с правилами
        agreed_to_rules = data.get('agreed_to_rules') == True

        if not agreed_to_rules:
            return jsonify({'error': 'Необходимо согласиться с правилами'}), 400

        # Для волонтеров требуется авторизация
        if reg_type == 'volunteer' and not current_user.is_authenticated:
            return jsonify(
                {'error': 'Для регистрации волонтером необходимо войти в систему', 'login_required': True}), 401

        # Проверяем обязательные поля
        required_fields = ['full_name', 'phone', 'birth_date']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'Поле {field} обязательно для заполнения'}), 400

        # Проверяем, не зарегистрирован ли уже пользователь (для авторизованных)
        if current_user.is_authenticated:
            existing = EventRegistration.query.filter_by(
                user_id=current_user.id,
                event_id=event_id
            ).first()

            if existing:
                return jsonify({'error': 'Вы уже зарегистрированы на это мероприятие'}), 400

        # Формируем данные для сохранения
        form_data = {
            'full_name': data.get('full_name', ''),
            'phone': data.get('phone', ''),
            'birth_date': data.get('birth_date', ''),
            'email': data.get('email', ''),
            'experience': data.get('experience', '')
        }

        user_id = current_user.id if current_user.is_authenticated else None

        registration = EventRegistration(
            user_id=user_id,
            event_id=event_id,
            registration_type=reg_type,
            form_data=json.dumps(form_data, ensure_ascii=False),
            agreed_to_rules=agreed_to_rules
        )

        db.session.add(registration)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Регистрация успешна'})

    except Exception as e:
        db.session.rollback()
        print(f"Error in registration: {e}")
        return jsonify({'error': 'Произошла ошибка при регистрации'}), 500


@app.route('/api/event/<int:event_id>/registration-data')
def get_registration_data(event_id):
    event = Event.query.get_or_404(event_id)

    if current_user.is_authenticated:
        user_data = {
            'full_name': current_user.get_full_name(),
            'phone': current_user.phone,
            'birth_date': current_user.birth_date.strftime('%Y-%m-%d') if current_user.birth_date else '',
            'email': current_user.email or '',
            'experience': current_user.experience or '',
            'is_authenticated': True
        }
    else:
        user_data = {
            'full_name': '',
            'phone': '',
            'birth_date': '',
            'email': '',
            'experience': '',
            'is_authenticated': False
        }

    return jsonify({
        'use_volunteer_template': event.use_volunteer_template,
        'use_participant_template': event.use_participant_template,
        'user_data': user_data
    })


# Админские маршруты
@app.route('/admin')
@login_required
def admin_dashboard():
    if current_user.role != 'admin':
        flash('Доступ запрещен', 'error')
        return redirect(url_for('index'))

    events = Event.query.order_by(Event.created_at.desc()).all()
    volunteers = User.query.filter_by(role='volunteer').all()

    # Получаем всех волонтеров (включая неавторизованных)
    all_volunteers = EventRegistration.query.filter_by(registration_type='volunteer').all()

    total_volunteers = User.query.filter_by(role='volunteer').count()
    total_events = Event.query.count()
    upcoming_events = Event.query.filter(Event.event_end_date >= date.today()).count()
    total_hours = db.session.query(db.func.sum(VolunteerHours.hours)).filter_by(status='approved').scalar() or 0

    stats = {
        'total_volunteers': total_volunteers,
        'total_events': total_events,
        'upcoming_events': upcoming_events,
        'total_hours': total_hours
    }

    return render_template('admin/dashboard.html',
                           events=events,
                           volunteers=volunteers,
                           all_volunteers=all_volunteers,
                           stats=stats)


@app.route('/admin/event/new', methods=['GET', 'POST'])
@login_required
def admin_event_new():
    if current_user.role != 'admin':
        return redirect(url_for('index'))

    if request.method == 'POST':
        use_volunteer = 'use_volunteer_template' in request.form
        use_participant = 'use_participant_template' in request.form

        # Получаем значение часов по умолчанию для волонтеров
        default_hours = request.form.get('default_hours', 0)
        try:
            default_hours = int(default_hours)
        except:
            default_hours = 0

        event = Event(
            title=request.form.get('title'),
            location=request.form.get('location', ''),
            registration_start_date=datetime.strptime(request.form.get('registration_start_date'), '%Y-%m-%d'),
            registration_end_date=datetime.strptime(request.form.get('registration_end_date'), '%Y-%m-%d'),
            event_start_date=datetime.strptime(request.form.get('event_start_date'), '%Y-%m-%d'),
            event_end_date=datetime.strptime(request.form.get('event_end_date'), '%Y-%m-%d'),
            description=request.form.get('description'),
            use_volunteer_template=use_volunteer,
            use_participant_template=use_participant,
            default_hours=default_hours,
            budget_checklist=request.form.get('budget_data', '[]'),
            created_by=current_user.id
        )

        if 'main_photo' in request.files:
            file = request.files['main_photo']
            if file.filename and allowed_file(file.filename):
                filename = secure_filename(f"event_main_{datetime.now().timestamp()}_{file.filename}")
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                event.main_photo = filename

        if 'additional_photos' in request.files:
            files = request.files.getlist('additional_photos')
            photo_list = []
            for file in files:
                if file.filename and allowed_file(file.filename):
                    filename = secure_filename(f"event_add_{datetime.now().timestamp()}_{file.filename}")
                    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    photo_list.append(filename)
            if photo_list:
                event.additional_photos = json.dumps(photo_list, ensure_ascii=False)

        material_names = request.form.getlist('material_name[]')
        material_files = request.files.getlist('material_file[]')

        materials = []
        for i, file in enumerate(material_files):
            if file and file.filename and allowed_file(file.filename):
                material_name = material_names[i] if i < len(material_names) and material_names[i] else file.filename
                filename = secure_filename(f"event_material_{datetime.now().timestamp()}_{i}_{file.filename}")
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                materials.append({
                    'name': material_name,
                    'file': filename
                })

        if materials:
            event.files = json.dumps(materials, ensure_ascii=False)

        db.session.add(event)
        db.session.commit()

        flash('Мероприятие создано', 'success')
        return redirect(url_for('admin_dashboard'))

    return render_template('admin/event_form.html')


@app.route('/admin/event/<int:event_id>')
@login_required
def admin_event_detail(event_id):
    if current_user.role != 'admin':
        return redirect(url_for('index'))

    event = Event.query.get_or_404(event_id)
    participants = EventRegistration.query.filter_by(event_id=event_id, registration_type='participant').all()
    volunteers = EventRegistration.query.filter_by(event_id=event_id, registration_type='volunteer').all()

    volunteer_data = []
    for reg in volunteers:
        hours = VolunteerHours.query.filter_by(
            user_id=reg.user_id,
            event_id=event_id
        ).first()
        volunteer_data.append({
            'registration': reg,
            'hours': hours
        })

    return render_template('admin/event_detail.html',
                           event=event,
                           participants=participants,
                           volunteers=volunteer_data)


@app.route('/admin/event/<int:event_id>/edit', methods=['GET', 'POST'])
@login_required
def admin_event_edit(event_id):
    if current_user.role != 'admin':
        return redirect(url_for('index'))

    event = Event.query.get_or_404(event_id)

    if request.method == 'POST':
        event.title = request.form.get('title')
        event.location = request.form.get('location', '')
        event.registration_start_date = datetime.strptime(request.form.get('registration_start_date'), '%Y-%m-%d')
        event.registration_end_date = datetime.strptime(request.form.get('registration_end_date'), '%Y-%m-%d')
        event.event_start_date = datetime.strptime(request.form.get('event_start_date'), '%Y-%m-%d')
        event.event_end_date = datetime.strptime(request.form.get('event_end_date'), '%Y-%m-%d')
        event.description = request.form.get('description')
        event.use_volunteer_template = 'use_volunteer_template' in request.form
        event.use_participant_template = 'use_participant_template' in request.form

        default_hours = request.form.get('default_hours', 0)
        try:
            event.default_hours = int(default_hours)
        except:
            event.default_hours = 0

        event.budget_checklist = request.form.get('budget_data', '[]')

        if 'main_photo' in request.files:
            file = request.files['main_photo']
            if file.filename and allowed_file(file.filename):
                if event.main_photo:
                    old_photo_path = os.path.join(app.config['UPLOAD_FOLDER'], event.main_photo)
                    if os.path.exists(old_photo_path):
                        os.remove(old_photo_path)

                filename = secure_filename(f"event_main_{datetime.now().timestamp()}_{file.filename}")
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                event.main_photo = filename

        if 'additional_photos' in request.files:
            files = request.files.getlist('additional_photos')
            existing_photos = []
            if event.additional_photos:
                try:
                    existing_photos = json.loads(event.additional_photos)
                except:
                    existing_photos = []

            for file in files:
                if file and file.filename and allowed_file(file.filename):
                    filename = secure_filename(f"event_add_{datetime.now().timestamp()}_{file.filename}")
                    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    existing_photos.append(filename)

            if existing_photos:
                event.additional_photos = json.dumps(existing_photos, ensure_ascii=False)

        material_names = request.form.getlist('material_name[]')
        material_files = request.files.getlist('material_file[]')

        existing_materials = []
        if event.files:
            try:
                existing_materials = json.loads(event.files)
            except:
                existing_materials = []

        for i, file in enumerate(material_files):
            if file and file.filename and allowed_file(file.filename):
                material_name = material_names[i] if i < len(material_names) and material_names[i] else file.filename
                filename = secure_filename(f"event_material_{datetime.now().timestamp()}_{i}_{file.filename}")
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                existing_materials.append({
                    'name': material_name,
                    'file': filename
                })

        if existing_materials:
            event.files = json.dumps(existing_materials, ensure_ascii=False)

        db.session.commit()
        flash('Мероприятие обновлено', 'success')
        return redirect(url_for('admin_event_detail', event_id=event.id))

    return render_template('admin/event_form.html', event=event)


@app.route('/admin/event/<int:event_id>/delete', methods=['POST'])
@login_required
def admin_event_delete(event_id):
    if current_user.role != 'admin':
        return jsonify({'error': 'Доступ запрещен'}), 403

    event = Event.query.get_or_404(event_id)

    if event.main_photo:
        photo_path = os.path.join(app.config['UPLOAD_FOLDER'], event.main_photo)
        if os.path.exists(photo_path):
            os.remove(photo_path)

    if event.additional_photos:
        photos = json.loads(event.additional_photos)
        for photo in photos:
            photo_path = os.path.join(app.config['UPLOAD_FOLDER'], photo)
            if os.path.exists(photo_path):
                os.remove(photo_path)

    if event.files:
        files = json.loads(event.files)
        for file_data in files:
            if isinstance(file_data, dict) and 'file' in file_data:
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_data['file'])
                if os.path.exists(file_path):
                    os.remove(file_path)

    EventRegistration.query.filter_by(event_id=event_id).delete()
    VolunteerHours.query.filter_by(event_id=event_id).delete()

    db.session.delete(event)
    db.session.commit()

    return jsonify({'success': True})


@app.route('/admin/event/<int:event_id>/remove-photo', methods=['POST'])
@login_required
def admin_event_remove_photo(event_id):
    if current_user.role != 'admin':
        return jsonify({'error': 'Доступ запрещен'}), 403

    data = request.get_json()
    photo_name = data.get('photo')

    event = Event.query.get_or_404(event_id)

    if event.additional_photos:
        try:
            photos = json.loads(event.additional_photos)
            if photo_name in photos:
                photos.remove(photo_name)
                photo_path = os.path.join(app.config['UPLOAD_FOLDER'], photo_name)
                if os.path.exists(photo_path):
                    os.remove(photo_path)
                event.additional_photos = json.dumps(photos, ensure_ascii=False)
                db.session.commit()
                return jsonify({'success': True})
        except:
            pass

    return jsonify({'error': 'Фото не найдено'}), 404


@app.route('/admin/event/<int:event_id>/remove-file', methods=['POST'])
@login_required
def admin_event_remove_file(event_id):
    if current_user.role != 'admin':
        return jsonify({'error': 'Доступ запрещен'}), 403

    data = request.get_json()
    file_name = data.get('file')

    event = Event.query.get_or_404(event_id)

    if event.files:
        try:
            files = json.loads(event.files)
            files = [f for f in files if f.get('file') != file_name]
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_name)
            if os.path.exists(file_path):
                os.remove(file_path)
            event.files = json.dumps(files, ensure_ascii=False)
            db.session.commit()
            return jsonify({'success': True})
        except:
            pass

    return jsonify({'error': 'Файл не найден'}), 404


@app.route('/admin/volunteer/hours', methods=['POST'])
@login_required
def admin_volunteer_hours():
    if current_user.role != 'admin':
        return jsonify({'error': 'Доступ запрещен'}), 403

    data = request.get_json()
    registration_id = data.get('registration_id')
    event_id = data.get('event_id')
    hours = int(data.get('hours', 0))
    fine = int(data.get('fine', 0))
    status = data.get('status', 'approved')
    comment = data.get('comment', '')

    registration = EventRegistration.query.get(registration_id)
    if not registration:
        return jsonify({'error': 'Регистрация не найдена'}), 404

    user_id = registration.user_id

    volunteer_hours = VolunteerHours.query.filter_by(
        user_id=user_id,
        event_id=event_id
    ).first()

    user = User.query.get(user_id) if user_id else None

    if volunteer_hours:
        # Если статус меняется, обновляем общее количество часов
        if user:
            # Вычитаем старые часы
            if volunteer_hours.status == 'approved':
                user.total_hours -= volunteer_hours.hours

            # Добавляем новые часы (с учетом штрафа)
            final_hours = hours - fine
            if status == 'approved':
                user.total_hours += final_hours

        volunteer_hours.hours = hours
        volunteer_hours.fine = fine
        volunteer_hours.status = status
        volunteer_hours.comment = comment
        volunteer_hours.approved_by = current_user.id
        volunteer_hours.approved_at = datetime.utcnow()
    else:
        final_hours = hours - fine
        volunteer_hours = VolunteerHours(
            user_id=user_id,
            event_id=event_id,
            hours=hours,
            fine=fine,
            status=status,
            comment=comment,
            approved_by=current_user.id,
            approved_at=datetime.utcnow()
        )
        db.session.add(volunteer_hours)

        if user and status == 'approved':
            user.total_hours += final_hours

    db.session.commit()

    return jsonify({'success': True})


@app.route('/admin/event/<int:event_id>/report')
@login_required
def generate_event_report(event_id):
    if current_user.role != 'admin':
        return redirect(url_for('index'))

    try:
        event = Event.query.get_or_404(event_id)
        participants = EventRegistration.query.filter_by(event_id=event_id, registration_type='participant').all()
        volunteers = EventRegistration.query.filter_by(event_id=event_id, registration_type='volunteer').all()

        doc = Document()

        title = doc.add_heading(f'Отчет по мероприятию: {event.title}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Основная информация', level=1)
        doc.add_paragraph(f'Место проведения: {event.location}')
        doc.add_paragraph(f'Даты проведения: {event.get_event_dates_string()}')
        doc.add_paragraph(f'Даты регистрации: {event.get_registration_dates_string()}')
        doc.add_paragraph(f'Описание: {event.description}')
        doc.add_paragraph(f'Регистрация волонтеров: {"Да" if event.use_volunteer_template else "Нет"}')
        doc.add_paragraph(f'Регистрация участников: {"Да" if event.use_participant_template else "Нет"}')
        doc.add_paragraph(f'Часы по умолчанию: {event.default_hours}')

        doc.add_heading('Статистика', level=1)
        doc.add_paragraph(f'Всего участников: {len(participants)}')
        doc.add_paragraph(f'Всего волонтеров: {len(volunteers)}')

        if event.budget_checklist:
            doc.add_heading('Бюджет мероприятия', level=1)
            try:
                budget = json.loads(event.budget_checklist)

                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Наименование'
                hdr_cells[1].text = 'Количество'
                hdr_cells[2].text = 'Цена'
                hdr_cells[3].text = 'Сумма'

                total = 0
                for item in budget:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item['name']
                    row_cells[1].text = str(item['quantity'])
                    row_cells[2].text = f"{item['price']} ₽"
                    sum_price = item['quantity'] * item['price']
                    row_cells[3].text = f"{sum_price} ₽"
                    total += sum_price

                doc.add_paragraph(f'Общий бюджет: {total} ₽')
            except:
                doc.add_paragraph('Ошибка загрузки данных бюджета')

        doc.add_heading('Список участников', level=1)
        if participants:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '№'
            hdr_cells[1].text = 'ФИО'
            hdr_cells[2].text = 'Телефон'
            hdr_cells[3].text = 'Email'
            hdr_cells[4].text = 'Дата регистрации'

            for i, reg in enumerate(participants, 1):
                form_data = json.loads(reg.form_data) if reg.form_data else {}
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = form_data.get('full_name', '') or (reg.user.get_full_name() if reg.user else '')
                row_cells[2].text = form_data.get('phone', '') or (reg.user.phone if reg.user else '')
                row_cells[3].text = form_data.get('email', '') or (reg.user.email if reg.user else '')
                row_cells[4].text = reg.registered_at.strftime('%d.%m.%Y %H:%M') if reg.registered_at else ''
        else:
            doc.add_paragraph('Нет зарегистрированных участников')

        doc.add_heading('Волонтеры и часы', level=1)
        if volunteers:
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '№'
            hdr_cells[1].text = 'ФИО'
            hdr_cells[2].text = 'Телефон'
            hdr_cells[3].text = 'Email'
            hdr_cells[4].text = 'Дата регистрации'
            hdr_cells[5].text = 'Часы'
            hdr_cells[6].text = 'Штраф'

            for i, reg in enumerate(volunteers, 1):
                form_data = json.loads(reg.form_data) if reg.form_data else {}
                hours_record = VolunteerHours.query.filter_by(user_id=reg.user_id, event_id=event_id).first()
                hours = hours_record.hours if hours_record else 0
                fine = hours_record.fine if hours_record else 0

                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = form_data.get('full_name', '') or (reg.user.get_full_name() if reg.user else '')
                row_cells[2].text = form_data.get('phone', '') or (reg.user.phone if reg.user else '')
                row_cells[3].text = form_data.get('email', '') or (reg.user.email if reg.user else '')
                row_cells[4].text = reg.registered_at.strftime('%d.%m.%Y %H:%M') if reg.registered_at else ''
                row_cells[5].text = str(hours)
                row_cells[6].text = str(fine)
        else:
            doc.add_paragraph('Нет зарегистрированных волонтеров')

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f'report_{event.title}_{datetime.now().strftime("%Y%m%d")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Error generating report: {e}")
        flash('Ошибка при генерации отчета', 'error')
        return redirect(url_for('admin_event_detail', event_id=event_id))


@app.route('/admin/event/<int:event_id>/export/participants/word')
@login_required
def export_event_participants_word(event_id):
    if current_user.role != 'admin':
        return redirect(url_for('index'))

    event = Event.query.get_or_404(event_id)
    participants = EventRegistration.query.filter_by(event_id=event_id, registration_type='participant').all()

    doc = Document()
    title = doc.add_heading(f'Участники мероприятия: {event.title}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Дата выгрузки: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph(f'Даты проведения: {event.get_event_dates_string()}')
    doc.add_paragraph()

    if participants:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'ФИО'
        hdr_cells[2].text = 'Телефон'
        hdr_cells[3].text = 'Email'
        hdr_cells[4].text = 'Дата регистрации'

        for i, reg in enumerate(participants, 1):
            form_data = json.loads(reg.form_data) if reg.form_data else {}
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = form_data.get('full_name', '') or (reg.user.get_full_name() if reg.user else '')
            row_cells[2].text = form_data.get('phone', '') or (reg.user.phone if reg.user else '')
            row_cells[3].text = form_data.get('email', '') or (reg.user.email if reg.user else '')
            row_cells[4].text = reg.registered_at.strftime('%d.%m.%Y %H:%M') if reg.registered_at else ''
    else:
        doc.add_paragraph('Нет зарегистрированных участников')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'participants_{event.title}_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/admin/event/<int:event_id>/export/volunteers/word')
@login_required
def export_event_volunteers_word(event_id):
    if current_user.role != 'admin':
        return redirect(url_for('index'))

    event = Event.query.get_or_404(event_id)
    volunteers = EventRegistration.query.filter_by(event_id=event_id, registration_type='volunteer').all()

    doc = Document()
    title = doc.add_heading(f'Волонтеры мероприятия: {event.title}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Дата выгрузки: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph(f'Даты проведения: {event.get_event_dates_string()}')
    doc.add_paragraph()

    if volunteers:
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'ФИО'
        hdr_cells[2].text = 'Телефон'
        hdr_cells[3].text = 'Email'
        hdr_cells[4].text = 'Дата регистрации'
        hdr_cells[5].text = 'Часы'
        hdr_cells[6].text = 'Штраф'

        for i, reg in enumerate(volunteers, 1):
            form_data = json.loads(reg.form_data) if reg.form_data else {}
            hours_record = VolunteerHours.query.filter_by(user_id=reg.user_id, event_id=event_id).first()
            hours = hours_record.hours if hours_record else 0
            fine = hours_record.fine if hours_record else 0

            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = form_data.get('full_name', '') or (reg.user.get_full_name() if reg.user else '')
            row_cells[2].text = form_data.get('phone', '') or (reg.user.phone if reg.user else '')
            row_cells[3].text = form_data.get('email', '') or (reg.user.email if reg.user else '')
            row_cells[4].text = reg.registered_at.strftime('%d.%m.%Y %H:%M') if reg.registered_at else ''
            row_cells[5].text = str(hours)
            row_cells[6].text = str(fine)
    else:
        doc.add_paragraph('Нет зарегистрированных волонтеров')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'volunteers_{event.title}_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/admin/volunteer/<int:user_id>')
@login_required
def admin_volunteer_detail(user_id):
    if current_user.role != 'admin':
        return jsonify({'error': 'Доступ запрещен'}), 403

    user = User.query.get_or_404(user_id)
    hours_history = VolunteerHours.query.filter_by(user_id=user_id).order_by(VolunteerHours.approved_at.desc()).all()
    events = EventRegistration.query.filter_by(user_id=user_id).order_by(EventRegistration.registered_at.desc()).all()

    data = {
        'id': user.id,
        'full_name': user.get_full_name(),
        'phone': user.phone,
        'email': user.email,
        'birth_date': user.birth_date.strftime('%d.%m.%Y') if user.birth_date else '',
        'experience': user.experience,
        'total_hours': user.total_hours,
        'avatar': url_for('static',
                          filename='uploads/' + user.avatar) if user.avatar and user.avatar != 'default-avatar.jpg' else url_for(
            'static', filename='images/default-avatar.jpg'),
        'hours_history': [
            {
                'event_title': h.event.title,
                'event_date': h.event.event_start_date.strftime('%d.%m.%Y'),
                'hours': h.hours,
                'fine': h.fine,
                'status': h.status,
                'comment': h.comment
            }
            for h in hours_history
        ]
    }

    return jsonify(data)


@app.route('/admin/export/volunteers/word')
@login_required
def export_all_volunteers_word():
    if current_user.role != 'admin':
        return redirect(url_for('index'))

    volunteers = EventRegistration.query.filter_by(registration_type='volunteer').order_by(
        EventRegistration.registered_at.desc()).all()

    doc = Document()
    title = doc.add_heading('Список всех волонтеров', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Дата выгрузки: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph()

    if volunteers:
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'ФИО'
        hdr_cells[2].text = 'Телефон'
        hdr_cells[3].text = 'Email'
        hdr_cells[4].text = 'Мероприятие'
        hdr_cells[5].text = 'Дата регистрации'
        hdr_cells[6].text = 'Часы'
        hdr_cells[7].text = 'Штраф'

        for i, reg in enumerate(volunteers, 1):
            form_data = json.loads(reg.form_data) if reg.form_data else {}
            hours_record = VolunteerHours.query.filter_by(user_id=reg.user_id, event_id=reg.event_id).first()
            hours = hours_record.hours if hours_record else 0
            fine = hours_record.fine if hours_record else 0

            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = form_data.get('full_name', '') or (reg.user.get_full_name() if reg.user else '')
            row_cells[2].text = form_data.get('phone', '') or (reg.user.phone if reg.user else '')
            row_cells[3].text = form_data.get('email', '') or (reg.user.email if reg.user else '')
            row_cells[4].text = reg.event.title if reg.event else ''
            row_cells[5].text = reg.registered_at.strftime('%d.%m.%Y %H:%M') if reg.registered_at else ''
            row_cells[6].text = str(hours)
            row_cells[7].text = str(fine)
    else:
        doc.add_paragraph('Нет зарегистрированных волонтеров')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'all_volunteers_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/admin/event/<int:event_id>/complete', methods=['POST'])
@login_required
def admin_event_complete(event_id):
    if current_user.role != 'admin':
        return jsonify({'error': 'Доступ запрещен'}), 403

    try:
        event = Event.query.get_or_404(event_id)

        # Получаем всех волонтеров мероприятия
        volunteers = EventRegistration.query.filter_by(event_id=event_id, registration_type='volunteer').all()

        # Для каждого волонтера создаем запись о часах, если её нет
        for reg in volunteers:
            if reg.user_id:  # Только для авторизованных волонтеров
                existing_hours = VolunteerHours.query.filter_by(
                    user_id=reg.user_id,
                    event_id=event_id
                ).first()

                if not existing_hours:
                    # Создаем запись с часами по умолчанию из мероприятия
                    volunteer_hours = VolunteerHours(
                        user_id=reg.user_id,
                        event_id=event_id,
                        hours=event.default_hours,
                        fine=0,
                        status='approved',
                        approved_by=current_user.id,
                        approved_at=datetime.utcnow()
                    )
                    db.session.add(volunteer_hours)

                    # Добавляем часы пользователю
                    user = User.query.get(reg.user_id)
                    if user:
                        user.total_hours += event.default_hours

        event.is_completed = True
        db.session.commit()

        return jsonify({'success': True})

    except Exception as e:
        db.session.rollback()
        print(f"Error completing event: {e}")
        return jsonify({'error': 'Ошибка при завершении мероприятия'}), 500


@app.route('/admin/event/<int:event_id>/revoke', methods=['POST'])
@login_required
def admin_event_revoke(event_id):
    if current_user.role != 'admin':
        return jsonify({'error': 'Доступ запрещен'}), 403

    try:
        event = Event.query.get_or_404(event_id)

        # Удаляем все записи о часах для этого мероприятия
        hours_records = VolunteerHours.query.filter_by(event_id=event_id).all()
        for hours in hours_records:
            if hours.user_id:
                # Вычитаем часы у пользователя
                user = User.query.get(hours.user_id)
                if user:
                    user.total_hours -= (hours.hours - hours.fine)
            db.session.delete(hours)

        event.is_completed = False
        db.session.commit()

        return jsonify({'success': True})

    except Exception as e:
        db.session.rollback()
        print(f"Error revoking event: {e}")
        return jsonify({'error': 'Ошибка при отзыве мероприятия'}), 500


@app.cli.command("init-db")
def init_db():
    db.create_all()

    admin_exists = False
    users = User.query.all()
    for u in users:
        if normalize_phone_for_db(u.username) == normalize_phone_for_db('89280319329'):
            admin_exists = True
            print(f'✓ Админ уже существует с номером: {u.username}')
            break

    if not admin_exists:
        formatted_admin_phone = format_phone('89280319329')
        admin = User(
            username=formatted_admin_phone,
            password=generate_password_hash('7777'),
            secret_word='admin',
            role='admin',
            first_name='Администратор',
            last_name='Системы',
            middle_name='',
            phone=formatted_admin_phone
        )
        db.session.add(admin)
        db.session.commit()
        print(f'✓ Админ создан с номером: {formatted_admin_phone}')

    volunteer_exists = False
    for u in users:
        if normalize_phone_for_db(u.username) == normalize_phone_for_db('89001112233'):
            volunteer_exists = True
            print(f'✓ Тестовый волонтер уже существует с номером: {u.username}')
            break

    if not volunteer_exists:
        formatted_volunteer_phone = format_phone('89001112233')
        volunteer = User(
            username=formatted_volunteer_phone,
            password=generate_password_hash('12345'),
            secret_word='volonter',
            role='volunteer',
            first_name='Иван',
            last_name='Иванов',
            middle_name='Иванович',
            phone=formatted_volunteer_phone,
            email='ivan@example.com',
            birth_date=datetime.strptime('2000-01-01', '%Y-%m-%d'),
            experience='Опыт волонтерства: 2 года'
        )
        db.session.add(volunteer)
        db.session.commit()
        print(f'✓ Тестовый волонтер создан с номером: {formatted_volunteer_phone}')

    print('База данных инициализирована!')


@app.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404


@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('500.html'), 500


if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        admin_exists = False
        users = User.query.all()
        for u in users:
            if normalize_phone_for_db(u.username) == normalize_phone_for_db('89280319329'):
                admin_exists = True
                print(f'✓ Админ уже существует с номером: {u.username}')
                break

        if not admin_exists:
            formatted_admin_phone = format_phone('89280319329')
            admin = User(
                username=formatted_admin_phone,
                password=generate_password_hash('7777'),
                secret_word='admin',
                role='admin',
                first_name='Администратор',
                last_name='Системы',
                phone=formatted_admin_phone
            )
            db.session.add(admin)
            db.session.commit()
            print(f'✓ Админ создан при запуске с номером: {formatted_admin_phone}')

        volunteer_exists = False
        for u in users:
            if normalize_phone_for_db(u.username) == normalize_phone_for_db('89001112233'):
                volunteer_exists = True
                print(f'✓ Тестовый волонтер уже существует с номером: {u.username}')
                break

        if not volunteer_exists:
            formatted_volunteer_phone = format_phone('89001112233')
            volunteer = User(
                username=formatted_volunteer_phone,
                password=generate_password_hash('12345'),
                secret_word='volonter',
                role='volunteer',
                first_name='Иван',
                last_name='Иванов',
                middle_name='Иванович',
                phone=formatted_volunteer_phone,
                email='ivan@example.com',
                birth_date=datetime.strptime('2000-01-01', '%Y-%m-%d'),
                experience='Опыт волонтерства: 2 года'
            )
            db.session.add(volunteer)
            db.session.commit()
            print(f'✓ Тестовый волонтер создан при запуске с номером: {formatted_volunteer_phone}')

    app.run(debug=True, host='0.0.0.0', port=5000)